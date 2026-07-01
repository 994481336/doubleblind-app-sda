from __future__ import annotations

from io import BytesIO
from html import escape

import pandas as pd

try:
    from docx import Document
except Exception:  # pragma: no cover
    Document = None


REPORT_TITLE = "飞行员双盲测试/复飞专项评估分析报告"
DISPLAY_COLUMN_NAMES = {
    "模板类型": "文件格式",
    "模板得分": "表内得分",
    "模板平均/合计得分": "平均/合计得分",
}


def display_df(df: pd.DataFrame) -> pd.DataFrame:
    return df.rename(columns=DISPLAY_COLUMN_NAMES) if not df.empty else df


def dataframe_to_html_table(df: pd.DataFrame, max_rows: int = 30) -> str:
    if df.empty:
        return "<p>暂无数据。</p>"
    return display_df(df).head(max_rows).to_html(index=False, border=0, classes="data-table", escape=True)


def build_html_report(
    pilot_scores: pd.DataFrame,
    company_df: pd.DataFrame,
    top_items: pd.DataFrame,
    quality_df: pd.DataFrame,
    subject_company_df: pd.DataFrame | None = None,
    risk_df: pd.DataFrame | None = None,
    comprehensive_df: pd.DataFrame | None = None,
) -> str:
    avg = pilot_scores["最终得分"].mean() if not pilot_scores.empty else 0
    html = f"""
<!doctype html>
<html lang="zh-CN">
<head>
<meta charset="utf-8">
<title>{escape(REPORT_TITLE)}</title>
<style>
body {{ font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", sans-serif; margin: 32px; color: #1f2937; }}
h1, h2 {{ color: #111827; }}
.metric {{ display: inline-block; min-width: 150px; padding: 12px 16px; margin: 6px 8px 6px 0; background: #f3f4f6; border-radius: 6px; }}
.metric b {{ font-size: 24px; display: block; margin-top: 4px; }}
.data-table {{ border-collapse: collapse; width: 100%; margin: 12px 0 24px; font-size: 13px; }}
.data-table th, .data-table td {{ border: 1px solid #d1d5db; padding: 8px; text-align: left; }}
.data-table th {{ background: #e5e7eb; }}
.note {{ color: #6b7280; }}
</style>
</head>
<body>
<h1>{escape(REPORT_TITLE)}</h1>
<p class="note">本报告基于用户本次上传文件即时生成，应用不保存上传文件或分析历史。</p>
<div class="metric">参评人数<b>{len(pilot_scores)}</b></div>
<div class="metric">平均分<b>{avg:.2f}</b></div>
<div class="metric">最高分<b>{pilot_scores['最终得分'].max() if not pilot_scores.empty else 0:.2f}</b></div>
<div class="metric">最低分<b>{pilot_scores['最终得分'].min() if not pilot_scores.empty else 0:.2f}</b></div>
<h2>格式质量检查</h2>
{dataframe_to_html_table(quality_df)}
<h2>单位统计</h2>
{dataframe_to_html_table(company_df)}
<h2>高频扣分项</h2>
{dataframe_to_html_table(top_items)}
<h2>单位科目平均失分</h2>
{dataframe_to_html_table(subject_company_df if subject_company_df is not None else pd.DataFrame())}
<h2>风险指标</h2>
{dataframe_to_html_table(risk_df if risk_df is not None else pd.DataFrame())}
<h2>综合考评</h2>
{dataframe_to_html_table(comprehensive_df if comprehensive_df is not None else pd.DataFrame())}
</body>
</html>
"""
    return html


def build_word_report(
    pilot_scores: pd.DataFrame,
    company_df: pd.DataFrame,
    top_items: pd.DataFrame,
    quality_df: pd.DataFrame,
    subject_company_df: pd.DataFrame | None = None,
    risk_df: pd.DataFrame | None = None,
    comprehensive_df: pd.DataFrame | None = None,
) -> BytesIO | None:
    if Document is None:
        return None

    doc = Document()
    doc.add_heading(REPORT_TITLE, level=0)
    doc.add_paragraph("本报告基于用户本次上传文件即时生成，应用不保存上传文件或分析历史。")
    if not pilot_scores.empty:
        doc.add_paragraph(f"参评人数：{len(pilot_scores)}")
        doc.add_paragraph(f"整体平均分：{pilot_scores['最终得分'].mean():.2f}")
        doc.add_paragraph(f"最高分：{pilot_scores['最终得分'].max():.2f}，最低分：{pilot_scores['最终得分'].min():.2f}")

    for title, df in [
        ("格式质量检查", quality_df),
        ("单位统计", company_df),
        ("高频扣分项", top_items),
        ("单位科目平均失分", subject_company_df if subject_company_df is not None else pd.DataFrame()),
        ("风险指标", risk_df if risk_df is not None else pd.DataFrame()),
        ("综合考评", comprehensive_df if comprehensive_df is not None else pd.DataFrame()),
    ]:
        doc.add_heading(title, level=1)
        if df.empty:
            doc.add_paragraph("暂无数据。")
            continue
        shown = display_df(df).head(12)
        table = doc.add_table(rows=1, cols=len(shown.columns))
        table.style = "Table Grid"
        for idx, col in enumerate(shown.columns):
            table.rows[0].cells[idx].text = str(col)
        for _, row in shown.iterrows():
            cells = table.add_row().cells
            for idx, col in enumerate(shown.columns):
                cells[idx].text = "" if pd.isna(row[col]) else str(row[col])

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio
